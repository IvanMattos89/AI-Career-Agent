from pathlib import Path
from shutil import copy2
from hashlib import sha256
from zipfile import BadZipFile, ZipFile
from dataclasses import dataclass
import re

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.database.sqlite_db import Database
from app.config import RESUMES_DIR


@dataclass(frozen=True)
class ResumeImportResult:
    resume_id: int
    destination: Path
    text: str


class ResumeService:
    MAX_FILE_SIZE = 10 * 1024 * 1024
    ALLOWED_EXTENSIONS = {".pdf", ".docx"}

    def __init__(self):
        self.resume_dir = RESUMES_DIR
        self.resume_dir.mkdir(parents=True, exist_ok=True)

        self.db = Database()

    def importar(self, arquivo):
        origem = Path(arquivo)
        self.validar_arquivo(origem)
        texto = self.ler(origem)
        if not texto.strip():
            raise ValueError("Não foi possível extrair texto do currículo. Use um PDF com texto selecionável ou um DOCX válido.")
        content_hash = sha256(texto.encode("utf-8", errors="ignore")).hexdigest()
        nome_seguro = re.sub(r"[^\w. -]", "_", origem.name, flags=re.UNICODE).strip(" .")
        nome_seguro = nome_seguro[-160:] or "curriculo"
        destino = self.resume_dir / f"{content_hash[:12]}_{nome_seguro}"
        if not destino.exists():
            try:
                copy2(origem, destino)
            except PermissionError as erro:
                raise PermissionError("Sem permissão para copiar o currículo para a pasta de dados.") from erro
            except OSError as erro:
                raise OSError("Não foi possível copiar o currículo selecionado.") from erro

        resume_id = self.db.salvar_curriculo(
            destino.name,
            str(destino),
            texto
        )

        return ResumeImportResult(resume_id=resume_id, destination=destino, text=texto)

    def validar_arquivo(self, arquivo):
        if not arquivo.exists() or not arquivo.is_file():
            raise FileNotFoundError("Arquivo de currículo não encontrado.")
        if arquivo.suffix.lower() not in self.ALLOWED_EXTENSIONS:
            raise ValueError("Formato não suportado. Selecione um arquivo PDF ou DOCX.")
        tamanho = arquivo.stat().st_size
        if tamanho == 0:
            raise ValueError("O arquivo selecionado está vazio.")
        if tamanho > self.MAX_FILE_SIZE:
            raise ValueError("O currículo excede o limite de 10 MB.")
        try:
            cabecalho = arquivo.read_bytes()[:8]
        except PermissionError as erro:
            raise PermissionError("Sem permissão para ler o currículo selecionado.") from erro
        if arquivo.suffix.lower() == ".pdf" and not cabecalho.startswith(b"%PDF-"):
            raise ValueError("O arquivo não possui uma estrutura PDF válida.")
        if arquivo.suffix.lower() == ".docx":
            try:
                with ZipFile(arquivo) as pacote:
                    if "word/document.xml" not in pacote.namelist():
                        raise ValueError("O arquivo não possui uma estrutura DOCX válida.")
            except BadZipFile as erro:
                raise ValueError("O arquivo DOCX está corrompido ou não é válido.") from erro

    def ler(self, arquivo):

        arquivo = Path(arquivo)

        if arquivo.suffix.lower() == ".docx":
            return self._ler_docx(arquivo)

        if arquivo.suffix.lower() == ".pdf":
            return self._ler_pdf(arquivo)

        raise ValueError("Formato não suportado.")

    def _ler_docx(self, arquivo):

        try:
            documento = Document(arquivo)
        except (BadZipFile, ValueError, KeyError) as erro:
            raise ValueError("Não foi possível abrir o DOCX. Verifique se o arquivo não está corrompido.") from erro
        partes = [p.text for p in documento.paragraphs if p.text.strip()]
        for tabela in documento.tables:
            for linha in tabela.rows:
                partes.append(" ".join(celula.text.strip() for celula in linha.cells if celula.text.strip()))
        return "\n".join(partes)

    def _ler_pdf(self, arquivo):

        try:
            reader = PdfReader(str(arquivo))
        except (PdfReadError, OSError) as erro:
            raise ValueError("Não foi possível abrir o PDF. Verifique se o arquivo não está corrompido.") from erro
        if reader.is_encrypted:
            raise ValueError("O PDF está protegido por senha. Remova a proteção antes de importar.")

        texto = []

        for pagina in reader.pages:

            conteudo = pagina.extract_text()

            if conteudo:
                texto.append(conteudo)

        return "\n".join(texto)
