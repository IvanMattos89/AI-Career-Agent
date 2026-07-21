from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.config import REPORTS_DIR


class ReportService:
    """Gera um relatório DOCX para um resultado de Job Match salvo ou recém-calculado."""

    def __init__(self):
        self.output_dir = REPORTS_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def exportar_job_match_docx(self, resultado):
        documento = Document()
        secao = documento.sections[0]
        secao.top_margin = secao.bottom_margin = Inches(0.75)
        estilos = documento.styles
        estilos["Normal"].font.name = "Aptos"
        estilos["Normal"].font.size = Pt(10.5)

        documento.add_heading("Relatório de Job Match", 0)
        documento.add_paragraph(f"Gerado em {datetime.now():%d/%m/%Y às %H:%M}")
        documento.add_paragraph(f"Currículo analisado: {resultado.get('curriculo', '-')}")
        documento.add_heading(f"Compatibilidade: {resultado.get('compatibilidade', 0)}%", 1)
        documento.add_paragraph(resultado.get("explicacao", ""))

        for titulo, chave in (
            ("Competências encontradas", "competencias_encontradas"),
            ("Competências a desenvolver", "competencias_faltantes"),
            ("Recomendações", "recomendacoes"),
        ):
            documento.add_heading(titulo, 1)
            itens = resultado.get(chave, [])
            if itens:
                for item in itens:
                    documento.add_paragraph(str(item), style="List Bullet")
            else:
                documento.add_paragraph("Nenhum item identificado.")

        documento.add_heading("Resumo", 1)
        documento.add_paragraph(resultado.get("resumo", "-"))
        destino = self.output_dir / f"job_match_{datetime.now():%Y%m%d_%H%M%S}.docx"
        documento.save(destino)
        return destino

    def exportar_pacote_candidatura_docx(self, pacote):
        documento = Document()
        secao = documento.sections[0]
        secao.top_margin = secao.bottom_margin = Inches(0.75)
        estilos = documento.styles
        estilos["Normal"].font.name = "Aptos"
        estilos["Normal"].font.size = Pt(10.5)
        documento.add_heading("Material de Candidatura", 0)
        documento.add_paragraph(f"Vaga: {pacote.get('vaga', '-')}")
        if pacote.get("empresa"):
            documento.add_paragraph(f"Empresa: {pacote['empresa']}")
        documento.add_heading("Carta de apresentação", 1)
        for paragrafo in pacote.get("carta", "").split("\n\n"):
            documento.add_paragraph(paragrafo)
        documento.add_heading("Resumo direcionado", 1)
        documento.add_paragraph(pacote.get("resumo_direcionado", ""))
        documento.add_heading("Palavras-chave para revisar", 1)
        for item in pacote.get("palavras_chave", []):
            documento.add_paragraph(item, style="List Bullet")
        documento.add_heading("Checklist antes de candidatar", 1)
        for item in pacote.get("checklist", []):
            documento.add_paragraph(item, style="List Bullet")
        destino = self.output_dir / f"candidatura_{datetime.now():%Y%m%d_%H%M%S}.docx"
        documento.save(destino)
        return destino

    def exportar_curriculo_adaptado_docx(self, dados):
        """Gera uma versão formal, revisável e direcionada da experiência já informada."""
        documento = Document()
        secao = documento.sections[0]
        # ABNT NBR 14724: esquerda/superior 3 cm; direita/inferior 2 cm.
        secao.top_margin = secao.left_margin = Cm(3)
        secao.bottom_margin = secao.right_margin = Cm(2)
        normal = documento.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(12)
        for nome, linhas in self._curriculo_blocos(dados):
            if nome == "__cabecalho__":
                nome_completo, contatos = linhas
                p = documento.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(nome_completo.upper())
                run.bold = True; run.font.name = "Arial"; run.font.size = Pt(15)
                p = documento.add_paragraph(contatos)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                continue
            p = documento.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(nome)
            run.bold = True; run.font.name = "Arial"; run.font.size = Pt(12)
            for linha in linhas:
                p = documento.add_paragraph(linha)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(3)
        destino = self.output_dir / f"curriculo_direcionado_{datetime.now():%Y%m%d_%H%M%S}.docx"
        documento.save(destino)
        return destino

    @staticmethod
    def _curriculo_blocos(dados):
        linhas = [linha.strip() for linha in dados["texto_original"].splitlines() if linha.strip()]
        cabecalho, secoes, atual = [], [], None
        titulos = {
            "OBJETIVO PROFISSIONAL", "SÍNTESE DE QUALIFICAÇÕES", "FORMACAO ACADÊMICA",
            "FORMAÇÃO ACADÊMICA", "HISTÓRICO PROFISSIONAL", "CONHECIMENTOS TÉCNICOS",
            "CURSOS COMPLEMENTARES",
        }
        for linha in linhas:
            if linha.upper() in titulos:
                atual = [linha.upper(), []]
                secoes.append(atual)
            elif atual is None:
                cabecalho.append(linha)
            else:
                atual[1].append(linha)
        nome = cabecalho[0] if cabecalho else "CURRÍCULO"
        contatos = " | ".join(cabecalho[1:])
        blocos = [("__cabecalho__", (nome, contatos))]
        blocos.append(("RESUMO PROFISSIONAL", [dados["resumo_direcionado"]]))
        competencias = dados["competencias_alinhadas"] or dados["palavras_revisar"][:8]
        if competencias:
            blocos.append(("COMPETÊNCIAS EM DESTAQUE", [" • ".join(competencias)]))
        for titulo, conteudo in secoes:
            if titulo == "OBJETIVO PROFISSIONAL":
                blocos.append((titulo, [f"Atuação direcionada à vaga de {dados['vaga']}.", *conteudo]))
            else:
                blocos.append((titulo, conteudo))
        return blocos

    def exportar_curriculo_adaptado_pdf(self, dados, destino):
        """Exporta o currículo em PDF com fontes Unicode e padrão formal A4."""
        # ReportLab usa as fontes-base PDF, que preservam acentos em português
        # sem depender da disponibilidade de fontes do Windows na máquina cliente.
        documento = SimpleDocTemplate(
            str(destino), pagesize=A4,
            leftMargin=3 * cm, rightMargin=2 * cm,
            topMargin=3 * cm, bottomMargin=2 * cm,
            title="Currículo direcionado",
        )
        estilos = getSampleStyleSheet()
        nome = ParagraphStyle(
            "NomeCurriculo", parent=estilos["Normal"], fontName="Helvetica-Bold",
            fontSize=15, leading=18, alignment=TA_CENTER, spaceAfter=3,
        )
        contato = ParagraphStyle(
            "ContatoCurriculo", parent=estilos["Normal"], fontName="Helvetica",
            fontSize=10, leading=13, alignment=TA_CENTER, spaceAfter=14,
        )
        titulo = ParagraphStyle(
            "TituloCurriculo", parent=estilos["Normal"], fontName="Helvetica-Bold",
            fontSize=12, leading=15, alignment=TA_LEFT, spaceBefore=8, spaceAfter=4,
        )
        corpo = ParagraphStyle(
            "CorpoCurriculo", parent=estilos["Normal"], fontName="Helvetica",
            fontSize=11, leading=16.5, alignment=TA_JUSTIFY, spaceAfter=3,
        )
        historia = []
        for secao, linhas in self._curriculo_blocos(dados):
            if secao == "__cabecalho__":
                nome_completo, contatos = linhas
                historia.append(Paragraph(self._escapar_pdf(nome_completo.upper()), nome))
                if contatos:
                    historia.append(Paragraph(self._escapar_pdf(contatos), contato))
                continue
            historia.append(Paragraph(self._escapar_pdf(secao), titulo))
            for linha in linhas:
                historia.append(Paragraph(self._escapar_pdf(linha), corpo))
            historia.append(Spacer(1, 2))
        documento.build(historia)

    @staticmethod
    def _escapar_pdf(texto):
        return (str(texto).replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;").replace("\n", "<br/>"))
